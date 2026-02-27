from docx import Document
from docx.shared import Pt

# Read the plain report file
src = 'INFORME_PROYECTO.docx'  # this currently contains plaintext
with open(src, 'r', encoding='utf-8') as f:
    text = f.read()

# Parse sections to build docx
lines = text.splitlines()

doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

i = 0
while i < len(lines):
    line = lines[i].strip()
    if line.startswith('# '):
        doc.add_heading(line[2:], level=1)
        i += 1
        continue
    if line.startswith('## '):
        doc.add_heading(line[3:], level=2)
        i += 1
        continue
    # Tables: detect markdown table header
    if line.startswith('| Tramo'):
        # read until blank line
        tbl_lines = []
        while i < len(lines) and lines[i].strip():
            tbl_lines.append(lines[i].rstrip())
            i += 1
        # parse markdown table
        header = tbl_lines[0].strip().strip('|').split('|')
        header = [h.strip() for h in header]
        rows = []
        for r in tbl_lines[2:]:
            cells = [c.strip() for c in r.strip().strip('|').split('|')]
            rows.append(cells)
        table = doc.add_table(rows=1+len(rows), cols=len(header))
        hdr_cells = table.rows[0].cells
        for idx, h in enumerate(header):
            hdr_cells[idx].text = h
        for r_idx, row in enumerate(rows, start=1):
            for c_idx, cell in enumerate(row):
                table.rows[r_idx].cells[c_idx].text = cell
        doc.add_paragraph()
        continue
    if line.startswith('| Parámetro'):
        # same parsing
        tbl_lines = []
        while i < len(lines) and lines[i].strip():
            tbl_lines.append(lines[i].rstrip())
            i += 1
        header = tbl_lines[0].strip().strip('|').split('|')
        header = [h.strip() for h in header]
        rows = []
        for r in tbl_lines[2:]:
            cells = [c.strip() for c in r.strip().strip('|').split('|')]
            rows.append(cells)
        table = doc.add_table(rows=1+len(rows), cols=len(header))
        hdr_cells = table.rows[0].cells
        for idx, h in enumerate(header):
            hdr_cells[idx].text = h
        for r_idx, row in enumerate(rows, start=1):
            for c_idx, cell in enumerate(row):
                table.rows[r_idx].cells[c_idx].text = cell
        doc.add_paragraph()
        continue
    if line.startswith('**Anexo A'):
        # add rest as preformatted until end
        doc.add_heading('Anexos', level=2)
        doc.add_paragraph('\n'.join(lines[i:]))
        break
    # normal paragraph
    if line:
        doc.add_paragraph(line)
    else:
        doc.add_paragraph()
    i += 1

out = 'INFORME_PROYECTO_formatted.docx'
doc.save(out)
print('Saved', out)
